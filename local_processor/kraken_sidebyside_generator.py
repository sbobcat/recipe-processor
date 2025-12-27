#!/usr/bin/env python3
"""
Side-by-Side Review Generator for Kraken OCR Results
Uses Kraken output from WSL to create Windows-compatible review document
"""

import json
import sys
from pathlib import Path
from typing import Dict, Any
import logging

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Missing python-docx. Install with: pip install python-docx")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)


class KrakenSideBySideGenerator:
    """Generate side-by-side review documents from Kraken OCR results."""

    def __init__(self, kraken_output_dir: str):
        """
        Initialize with Kraken output directory.

        Args:
            kraken_output_dir: Path to directory containing Kraken results
        """
        self.output_dir = Path(kraken_output_dir)
        self.results_file = self.output_dir / "processing_results.json"
        self.images_dir = self.output_dir / "page_images"

        # Validate inputs
        if not self.output_dir.exists():
            raise FileNotFoundError(
                f"Kraken output directory not found: {self.output_dir}"
            )

        if not self.results_file.exists():
            raise FileNotFoundError(
                f"Kraken results file not found: {self.results_file}"
            )

    def load_kraken_results(self) -> Dict[str, Any]:
        """Load Kraken processing results from JSON file."""
        with open(self.results_file, 'r', encoding='utf-8') as f:
            return json.load(f)

    def create_review_document(self, output_path: str = None) -> Path:
        """
        Create Word document with side-by-side image and OCR text comparison.

        Args:
            output_path: Optional custom output path

        Returns:
            Path to created document
        """
        # Load results
        results = self.load_kraken_results()

        # Determine output path
        if output_path:
            doc_path = Path(output_path)
        else:
            doc_path = (self.output_dir.parent /
                        "Anns_Recipes_Kraken_Review.docx")

        logger.info(f"Creating review document: {doc_path}")

        # Create Word document
        doc = Document()

        # Add title and instructions
        doc.add_heading("Ann's Recipes — Kraken OCR Review", level=1)

        # Add processing info
        info_para = doc.add_paragraph()
        info_para.add_run("Processing Information:").bold = True
        info_para.add_run(f"\n• Total pages: {results['total_pages']}")
        info_para.add_run(
            f"\n• Successfully processed: {results['successful_pages']}"
        )
        info_para.add_run(
            f"\n• Segmentation model: "
            f"{results.get('segmentation_model', 'Unknown')}"
        )
        info_para.add_run(
            f"\n• Recognition model: "
            f"{results.get('recognition_model', 'Unknown')}"
        )
        info_para.add_run(
            f"\n• Source PDF: {Path(results['pdf_file']).name}"
        )

        # Add instructions
        instructions = doc.add_paragraph()
        instructions.add_run("\nInstructions: ").bold = True
        instructions.add_run(
            "Compare the original scan (left) with the Kraken OCR text "
            "(right). Edit the text directly in this document to correct "
            "any errors. Kraken is optimized for handwritten text, so "
            "results should be quite good!\n"
        )

        # Process each page
        successful_pages = [p for p in results['pages'] if p['success']]
        failed_pages = [p for p in results['pages'] if not p['success']]

        if failed_pages:
            doc.add_paragraph(
                f"⚠️ Note: {len(failed_pages)} pages failed OCR processing "
                f"and are not included."
            )

        for page_data in successful_pages:
            page_num = page_data['page_number']

            # Map to actual file locations based on your Kraken processor
            # Images are in: kraken_output/page_images/page-001.png, etc.
            image_filename = f"page-{page_num:03d}.png"
            image_path = self.images_dir / image_filename

            # Text files are in: kraken_output/page_001_text.txt, etc.
            text_filename = f"page_{page_num:03d}_text.txt"
            text_file_path = self.output_dir / text_filename

            # Read OCR text from the actual text file
            if text_file_path.exists():
                with open(text_file_path, 'r', encoding='utf-8') as f:
                    ocr_text = f.read().strip()
            else:
                ocr_text = page_data.get('text', '')

            logger.info(f"Adding page {page_num} to document")
            logger.info(f"  Image: {image_path}")
            logger.info(f"  Text file: {text_file_path}")

            # Add page header
            heading = doc.add_paragraph()
            heading.add_run(f"Page {page_num}").bold = True

            # Create two-column table
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Set column widths
            table.columns[0].width = Inches(4.0)  # Image column
            table.columns[1].width = Inches(4.0)  # Text column

            left_cell = table.cell(0, 0)
            right_cell = table.cell(0, 1)

            # Left column: Image
            left_para = left_cell.paragraphs[0]
            if image_path.exists():
                try:
                    left_run = left_para.add_run()
                    left_run.add_picture(str(image_path), width=Inches(3.8))
                except Exception as e:
                    logger.warning(
                        f"Failed to add image for page {page_num}: {e}"
                    )
                    left_para.add_run(f"[Image error: {e}]")
            else:
                left_para.add_run(f"[Image not found: {image_path}]")
                logger.warning(f"Image not found: {image_path}")

            # Right column: Kraken OCR text
            right_para = right_cell.paragraphs[0]
            if ocr_text.strip():
                right_para.text = ocr_text
            else:
                right_para.text = "[No text detected by Kraken OCR]"

            # Format text for readability
            for run in right_para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

            # Add page break (except for last page)
            if page_num < len(successful_pages):
                doc.add_page_break()

        # Add summary section
        if failed_pages:
            doc.add_page_break()
            doc.add_heading("Failed Pages", level=2)

            for page_data in failed_pages:
                para = doc.add_paragraph()
                para.add_run(f"Page {page_data['page_number']}: ").bold = True
                para.add_run(page_data.get('error', 'Unknown error'))

        # Save document
        doc.save(str(doc_path))
        logger.info(f"✓ Review document created: {doc_path}")

        # Show file info
        file_size_mb = doc_path.stat().st_size / (1024 * 1024)
        logger.info(f"File size: {file_size_mb:.2f} MB")

        return doc_path


def main():
    """Main entry point."""
    # Configuration - adjust path as needed
    kraken_output_dir = (
        r"C:\Users\steph\OneDrive\Documents\Scanned Documents"
        r"\annsrecipes\kraken_output"
    )

    try:
        generator = KrakenSideBySideGenerator(kraken_output_dir)
        doc_path = generator.create_review_document()

        print("\n🎉 Kraken review document created successfully!")
        print(f"📄 Open: {doc_path}")
        print("\n💡 Tips:")
        print("- Kraken is optimized for handwritten text")
        print("- Edit the OCR text directly in Word")
        print("- Use Track Changes to see your edits")
        print("- Save frequently as you work")

    except FileNotFoundError as e:
        logger.error(f"File not found: {e}")
        print("\n❓ Make sure you've run the Kraken processing first:")
        print("1. Set up WSL Ubuntu")
        print("2. Install Kraken in WSL")
        print("3. Run process_recipes_kraken.py in WSL")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
