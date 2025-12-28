#!/usr/bin/env python3
"""
Side-by-Side Review Generator for AWS Textract OCR Results
Creates Windows-compatible review documents from AWS Textract output
"""

import json
import sys
from pathlib import Path
from typing import Dict, Any, List
import logging
import tempfile
import io

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import RGBColor
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:
    print("Missing python-docx. Install with: pip install python-docx")
    sys.exit(1)

try:
    import fitz  # PyMuPDF
    from PIL import Image
except ImportError:
    print("Missing PyMuPDF or PIL. Install with: pip install PyMuPDF Pillow")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)


class AWSTextractSideBySideGenerator:
    """Generate side-by-side review documents from AWS Textract OCR results."""

    def __init__(self, aws_output_dir: str):
        """
        Initialize with AWS Textract output directory.

        Args:
            aws_output_dir: Path to directory containing AWS Textract results
        """
        self.output_dir = Path(aws_output_dir)
        
        # Find the summary JSON file (should be the only JSON file)
        json_files = list(self.output_dir.glob("*_summary.json"))
        if not json_files:
            raise FileNotFoundError(
                f"No summary JSON file found in: {self.output_dir}"
            )
        
        self.results_file = json_files[0]
        
        # Validate inputs
        if not self.output_dir.exists():
            raise FileNotFoundError(
                f"AWS output directory not found: {self.output_dir}"
            )

        if not self.results_file.exists():
            raise FileNotFoundError(
                f"AWS results file not found: {self.results_file}"
            )

    def load_aws_results(self) -> Dict[str, Any]:
        """Load AWS Textract processing results from JSON file."""
        with open(self.results_file, 'r', encoding='utf-8') as f:
            return json.load(f)

    def extract_pdf_images(self, pdf_path: Path) -> List[bytes]:
        """
        Extract images from PDF pages for inclusion in review document.
        
        Args:
            pdf_path: Path to the source PDF file
            
        Returns:
            List of PNG image data as bytes
        """
        if not pdf_path.exists():
            logger.warning(f"Source PDF not found: {pdf_path}")
            return []
            
        try:
            doc = fitz.open(str(pdf_path))
            images = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # Use same resolution as AWS processor for consistency
                pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))  # 216 DPI
                img_data = pix.tobytes("png")
                images.append(img_data)
            
            doc.close()
            logger.info(f"Extracted {len(images)} page images from PDF")
            return images
            
        except Exception as e:
            logger.error(f"Failed to extract images from PDF: {e}")
            return []

    def save_temp_image(self, image_data: bytes, page_num: int, temp_dir: Path) -> Path:
        """
        Save image data to a temporary file for Word document inclusion.
        
        Args:
            image_data: PNG image data as bytes
            page_num: Page number for filename
            temp_dir: Temporary directory path
            
        Returns:
            Path to saved image file
        """
        image_path = temp_dir / f"page_{page_num:03d}.png"
        
        try:
            with open(image_path, 'wb') as f:
                f.write(image_data)
            return image_path
        except Exception as e:
            logger.error(f"Failed to save temp image for page {page_num}: {e}")
            return None

    def create_review_document(self, output_path: str = None) -> Path:
        """
        Create Word document with side-by-side image and OCR text comparison.

        Args:
            output_path: Optional custom output path

        Returns:
            Path to created document
        """
        # Load results
        results = self.load_aws_results()

        # Determine output path
        if output_path:
            doc_path = Path(output_path)
        else:
            pdf_name = results.get('pdf_name', 'Unknown')
            base_name = Path(pdf_name).stem
            doc_path = (self.output_dir.parent /
                        f"{base_name}_AWS_Textract_Review.docx")

        logger.info(f"Creating AWS Textract review document: {doc_path}")

        # Find the source PDF file
        pdf_name = results.get('pdf_name', 'Unknown')
        # Look for the PDF in common locations
        possible_pdf_paths = [
            self.output_dir.parent / pdf_name,  # Same directory as output
            self.output_dir.parent.parent / pdf_name,  # Parent directory
            Path(pdf_name),  # Absolute path if provided
        ]
        
        source_pdf_path = None
        for pdf_path in possible_pdf_paths:
            if pdf_path.exists():
                source_pdf_path = pdf_path
                break
        
        # Extract images from source PDF
        page_images = []
        temp_dir = None
        if source_pdf_path:
            logger.info(f"Found source PDF: {source_pdf_path}")
            page_images = self.extract_pdf_images(source_pdf_path)
            
            # Create temporary directory for image files
            if page_images:
                temp_dir = Path(tempfile.mkdtemp())
                logger.info(f"Created temporary directory: {temp_dir}")
        else:
            logger.warning(f"Source PDF not found. Tried: {[str(p) for p in possible_pdf_paths]}")

        # Create Word document
        doc = Document()

        # Add title and instructions
        doc.add_heading("AWS Textract OCR Review", level=1)

        # Add processing info
        info_para = doc.add_paragraph()
        info_para.add_run("Processing Information:").bold = True
        info_para.add_run(f"\n• Total pages: {results['total_pages']}")
        
        # Calculate successful pages
        successful_pages = [p for p in results['pages'] if not p.get('has_error', False)]
        failed_pages = [p for p in results['pages'] if p.get('has_error', False)]
        
        info_para.add_run(f"\n• Successfully processed: {len(successful_pages)}")
        info_para.add_run(f"\n• Processing method: AWS Textract")
        info_para.add_run(f"\n• Source PDF: {results.get('pdf_name', 'Unknown')}")
        
        # Calculate average confidence
        if successful_pages:
            avg_confidence = sum(p.get('confidence', 0) for p in successful_pages) / len(successful_pages)
            info_para.add_run(f"\n• Average confidence: {avg_confidence:.1f}%")

        # Add instructions
        instructions = doc.add_paragraph()
        instructions.add_run("\nInstructions: ").bold = True
        instructions.add_run(
            "Compare the original scan (left) with the AWS Textract OCR text "
            "(right). Edit the text directly in this document to correct "
            "any errors. Words highlighted in yellow have low confidence "
            "scores and may need extra attention.\n"
        )

        # Add confidence threshold info
        confidence_info = doc.add_paragraph()
        confidence_info.add_run("Confidence Scoring: ").bold = True
        confidence_info.add_run(
            "AWS Textract provides confidence scores for each word. "
            "Words with confidence below 80% are highlighted for review. "
            "Higher confidence scores generally indicate more accurate text extraction.\n"
        )

        if failed_pages:
            doc.add_paragraph(
                f"⚠️ Note: {len(failed_pages)} pages had processing errors "
                f"and are not included."
            )

        # Process each successful page
        for page_data in successful_pages:
            page_num = page_data['page_number']

            # Read OCR text from the individual page file
            text_filename = f"page_{page_num:03d}_ocr.txt"
            text_file_path = self.output_dir / text_filename

            if text_file_path.exists():
                with open(text_file_path, 'r', encoding='utf-8') as f:
                    full_text = f.read()
                    # Extract just the OCR text (between the separator lines)
                    lines = full_text.split('\n')
                    ocr_text = ""
                    in_text_section = False
                    for line in lines:
                        if line.startswith("="):
                            if in_text_section:
                                break  # End of text section
                            else:
                                in_text_section = True  # Start of text section
                        elif in_text_section and not line.startswith("LOW CONFIDENCE"):
                            ocr_text += line + "\n"
                    ocr_text = ocr_text.strip()
            else:
                ocr_text = page_data.get('text', '')

            logger.info(f"Adding page {page_num} to document")
            logger.info(f"  Text file: {text_file_path}")
            logger.info(f"  Confidence: {page_data.get('confidence', 0):.1f}%")

            # Add page header
            heading = doc.add_paragraph()
            heading_run = heading.add_run(f"Page {page_num}")
            heading_run.bold = True
            
            # Add confidence info for this page
            confidence = page_data.get('confidence', 0)
            conf_run = heading.add_run(f" (Confidence: {confidence:.1f}%)")
            if confidence < 70:
                conf_run.font.color.rgb = RGBColor(255, 0, 0)  # Red for low confidence
            elif confidence < 85:
                conf_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange for medium confidence
            else:
                conf_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for high confidence

            # Create two-column table
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Set column widths
            table.columns[0].width = Inches(4.0)  # Image column
            table.columns[1].width = Inches(4.0)  # Text column

            left_cell = table.cell(0, 0)
            right_cell = table.cell(0, 1)

            # Left column: Actual page image or placeholder
            left_para = left_cell.paragraphs[0]
            
            # Try to add the actual page image
            image_added = False
            if page_images and len(page_images) >= page_num and temp_dir:
                try:
                    # Save image to temporary file
                    image_data = page_images[page_num - 1]  # Convert to 0-based index
                    temp_image_path = self.save_temp_image(image_data, page_num, temp_dir)
                    
                    if temp_image_path and temp_image_path.exists():
                        # Add image to document
                        left_run = left_para.add_run()
                        left_run.add_picture(str(temp_image_path), width=Inches(3.8))
                        image_added = True
                        logger.info(f"  Added image for page {page_num}")
                    
                except Exception as e:
                    logger.warning(f"Failed to add image for page {page_num}: {e}")
            
            # If image couldn't be added, show instructions
            if not image_added:
                left_para.add_run("[Original PDF page image would appear here]")
                left_para.add_run(
                    f"\n\nTo view the original image:\n"
                    f"1. Open the source PDF: {results.get('pdf_name', 'Unknown')}\n"
                    f"2. Navigate to page {page_num}\n"
                    f"3. Compare with the OCR text on the right"
                )

            # Right column: AWS Textract OCR text with confidence highlighting
            right_para = right_cell.paragraphs[0]
            if ocr_text.strip():
                # For now, add the text normally
                # In a more advanced version, we could parse individual word confidences
                # and highlight low-confidence words
                right_para.text = ocr_text
                
                # Highlight the entire text if page confidence is low
                if confidence < 80:
                    for run in right_para.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                right_para.text = "[No text detected by AWS Textract]"

            # Format text for readability
            for run in right_para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

            # Add low confidence words info if available
            if text_file_path.exists():
                with open(text_file_path, 'r', encoding='utf-8') as f:
                    full_text = f.read()
                    if "LOW CONFIDENCE WORDS" in full_text:
                        # Extract low confidence words section
                        low_conf_section = full_text.split("LOW CONFIDENCE WORDS")[1]
                        if "None - all words have good confidence!" not in low_conf_section:
                            # Add a note about low confidence words
                            right_para.add_run("\n\n--- Low Confidence Words ---\n")
                            low_conf_run = right_para.add_run(
                                "⚠️ Some words in this page have low confidence scores. "
                                "Please review carefully against the original image."
                            )
                            low_conf_run.font.color.rgb = RGBColor(255, 140, 0)  # Orange
                            low_conf_run.italic = True

            # Add page break (except for last page)
            if page_num < len(successful_pages):
                doc.add_page_break()

        # Add summary section for failed pages
        if failed_pages:
            doc.add_page_break()
            doc.add_heading("Processing Errors", level=2)

            error_para = doc.add_paragraph()
            error_para.add_run("The following pages encountered processing errors:\n").bold = True
            
            for page_data in failed_pages:
                error_para.add_run(f"• Page {page_data['page_number']}: ")
                error_para.add_run("Processing error occurred\n")

        # Add statistics summary
        doc.add_page_break()
        doc.add_heading("Processing Statistics", level=2)
        
        stats_para = doc.add_paragraph()
        stats_para.add_run("Summary Statistics:\n").bold = True
        stats_para.add_run(f"• Total pages processed: {results['total_pages']}\n")
        stats_para.add_run(f"• Successful pages: {len(successful_pages)}\n")
        stats_para.add_run(f"• Failed pages: {len(failed_pages)}\n")
        
        if successful_pages:
            avg_confidence = sum(p.get('confidence', 0) for p in successful_pages) / len(successful_pages)
            stats_para.add_run(f"• Average confidence: {avg_confidence:.1f}%\n")
            
            # Count pages by confidence level
            high_conf = len([p for p in successful_pages if p.get('confidence', 0) >= 85])
            med_conf = len([p for p in successful_pages if 70 <= p.get('confidence', 0) < 85])
            low_conf = len([p for p in successful_pages if p.get('confidence', 0) < 70])
            
            stats_para.add_run(f"• High confidence pages (≥85%): {high_conf}\n")
            stats_para.add_run(f"• Medium confidence pages (70-84%): {med_conf}\n")
            stats_para.add_run(f"• Low confidence pages (<70%): {low_conf}\n")

        # Save document
        doc.save(str(doc_path))
        logger.info(f"✓ AWS Textract review document created: {doc_path}")

        # Clean up temporary directory
        if temp_dir and temp_dir.exists():
            try:
                import shutil
                shutil.rmtree(temp_dir)
                logger.info(f"Cleaned up temporary directory: {temp_dir}")
            except Exception as e:
                logger.warning(f"Failed to clean up temporary directory: {e}")

        # Show file info
        file_size_mb = doc_path.stat().st_size / (1024 * 1024)
        logger.info(f"File size: {file_size_mb:.2f} MB")

        return doc_path


def main():
    """Main entry point."""
    # Configuration - adjust path as needed
    aws_output_dir = (
        r"C:\Code\pers\recipe-processor\test-data\aws_textract_output"
    )

    try:
        generator = AWSTextractSideBySideGenerator(aws_output_dir)
        doc_path = generator.create_review_document()

        print("\n🎉 AWS Textract review document created successfully!")
        print(f"📄 Open: {doc_path}")
        print("\n💡 Tips:")
        print("- AWS Textract provides confidence scores for each word")
        print("- Yellow highlighting indicates low confidence text")
        print("- Edit the OCR text directly in Word")
        print("- Use Track Changes to see your edits")
        print("- Compare with the original PDF for accuracy")
        print("- Save frequently as you work")

    except FileNotFoundError as e:
        logger.error(f"File not found: {e}")
        print("\n❓ Make sure you've run the AWS Textract processing first:")
        print("1. Set up AWS credentials (aws configure)")
        print("2. Run kraken_alternative_aws.py")
        print("3. Check that the output directory contains the summary JSON file")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()