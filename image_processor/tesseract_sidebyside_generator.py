#!/usr/bin/env python3
"""
Side-by-Side Review Generator for Tesseract OCR Results
Creates Windows-compatible review documents from Tesseract image processing output
"""

import json
import sys
from pathlib import Path
from typing import Dict, Any, List
import logging

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import RGBColor
except ImportError:
    print("Missing python-docx. Install with: pip install python-docx")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)


class TesseractSideBySideGenerator:
    """Generate side-by-side review documents from Tesseract OCR results."""

    def __init__(self, tesseract_output_dir: str):
        """
        Initialize with Tesseract output directory.

        Args:
            tesseract_output_dir: Path to directory containing Tesseract results
        """
        self.output_dir = Path(tesseract_output_dir)
        
        # Find the summary JSON file
        json_files = list(self.output_dir.glob("*_tesseract_summary.json"))
        if not json_files:
            raise FileNotFoundError(
                f"No Tesseract summary JSON file found in: {self.output_dir}"
            )
        
        self.results_file = json_files[0]
        
        # Validate inputs
        if not self.output_dir.exists():
            raise FileNotFoundError(
                f"Tesseract output directory not found: {self.output_dir}"
            )

        if not self.results_file.exists():
            raise FileNotFoundError(
                f"Tesseract results file not found: {self.results_file}"
            )

    def load_tesseract_results(self) -> Dict[str, Any]:
        """Load Tesseract processing results from JSON file."""
        with open(self.results_file, 'r', encoding='utf-8') as f:
            return json.load(f)

    def create_review_document(self, output_path: str = None) -> Path:
        """
        Create Word document with side-by-side image and OCR text comparison.

        Args:
            output_path: Optional custom output path

        Returns:
            Path to created document
        """
        # Load results
        results = self.load_tesseract_results()

        # Determine output path
        if output_path:
            doc_path = Path(output_path)
        else:
            folder_name = Path(results.get('folder_name', 'Unknown')).name
            doc_path = (self.output_dir.parent /
                        f"{folder_name}_Tesseract_Review.docx")

        logger.info(f"Creating Tesseract review document: {doc_path}")

        # Create Word document
        doc = Document()

        # Add title and instructions
        doc.add_heading("Tesseract OCR Review", level=1)

        # Add processing info
        info_para = doc.add_paragraph()
        info_para.add_run("Processing Information:").bold = True
        info_para.add_run(f"\n• Total images: {results['total_images']}")
        
        # Calculate successful and failed images
        successful_images = [img for img in results['images'] if img.get('success', False)]
        failed_images = [img for img in results['images'] if not img.get('success', False)]
        
        info_para.add_run(f"\n• Successfully processed: {len(successful_images)}")
        info_para.add_run(f"\n• Processing method: Tesseract OCR")
        info_para.add_run(f"\n• Source folder: {results.get('folder_name', 'Unknown')}")
        
        # Calculate average confidence if available
        images_with_confidence = [img for img in successful_images if 'confidence' in img]
        if images_with_confidence:
            avg_confidence = sum(img['confidence'] for img in images_with_confidence) / len(images_with_confidence)
            info_para.add_run(f"\n• Average confidence: {avg_confidence:.1f}%")

        # Add instructions
        instructions = doc.add_paragraph()
        instructions.add_run("\nInstructions: ").bold = True
        instructions.add_run(
            "Compare the original image (left) with the Tesseract OCR text "
            "(right). Edit the text directly in this document to correct "
            "any errors. Tesseract is optimized for printed text and works "
            "best with clear, high-resolution images.\n"
        )

        if failed_images:
            doc.add_paragraph(
                f"⚠️ Note: {len(failed_images)} images had processing errors "
                f"and are not included."
            )

        # Process each successful image
        for image_data in successful_images:
            image_num = image_data['image_number']
            image_file = Path(image_data['image_file'])

            # Read OCR text from the text file
            text_file_path = Path(image_data.get('text_file', ''))
            
            if text_file_path.exists():
                with open(text_file_path, 'r', encoding='utf-8') as f:
                    full_text = f.read()
                    # Extract just the OCR text (between the separator lines)
                    lines = full_text.split('\n')
                    ocr_text = ""
                    in_text_section = False
                    for line in lines:
                        if line.startswith("="):
                            if in_text_section:
                                break  # End of text section
                            else:
                                in_text_section = True  # Start of text section
                        elif in_text_section and not line.startswith("LOW CONFIDENCE"):
                            ocr_text += line + "\n"
                    ocr_text = ocr_text.strip()
            else:
                ocr_text = image_data.get('text', '')

            logger.info(f"Adding image {image_num} to document")
            logger.info(f"  Image file: {image_file}")
            logger.info(f"  Text file: {text_file_path}")
            
            confidence = image_data.get('confidence', 0)
            if confidence > 0:
                logger.info(f"  Confidence: {confidence:.1f}%")

            # Add image header
            heading = doc.add_paragraph()
            heading_run = heading.add_run(f"Image {image_num}")
            heading_run.bold = True
            
            # Add filename
            filename_run = heading.add_run(f" ({image_file.name})")
            filename_run.font.size = Pt(9)
            filename_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add confidence info if available
            if confidence > 0:
                conf_run = heading.add_run(f" - Confidence: {confidence:.1f}%")
                if confidence < 70:
                    conf_run.font.color.rgb = RGBColor(255, 0, 0)  # Red for low confidence
                elif confidence < 85:
                    conf_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange for medium confidence
                else:
                    conf_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for high confidence

            # Create two-column table
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Set column widths
            table.columns[0].width = Inches(4.0)  # Image column
            table.columns[1].width = Inches(4.0)  # Text column

            left_cell = table.cell(0, 0)
            right_cell = table.cell(0, 1)

            # Left column: Original image
            left_para = left_cell.paragraphs[0]
            
            if image_file.exists():
                try:
                    # Calculate appropriate image width to maintain aspect ratio
                    from PIL import Image
                    with Image.open(image_file) as img:
                        width, height = img.size
                        aspect_ratio = height / width
                        
                        # Target width is 3.8 inches
                        target_width = 3.8
                        
                        # If image is very tall, limit height instead
                        if aspect_ratio > 2:  # Very tall image
                            target_height = 7.0  # Max height in inches
                            target_width = target_height / aspect_ratio
                    
                    left_run = left_para.add_run()
                    left_run.add_picture(str(image_file), width=Inches(target_width))
                    logger.info(f"  Added image with width {target_width:.2f} inches")
                    
                except Exception as e:
                    logger.warning(f"Failed to add image for image {image_num}: {e}")
                    left_para.add_run(f"[Image error: {e}]")
            else:
                left_para.add_run(f"[Image not found: {image_file}]")
                logger.warning(f"Image not found: {image_file}")

            # Right column: Tesseract OCR text
            right_para = right_cell.paragraphs[0]
            if ocr_text.strip():
                right_para.text = ocr_text
            else:
                right_para.text = "[No text detected by Tesseract OCR]"

            # Format text for readability
            for run in right_para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

            # Add page break (except for last image)
            if image_num < len(successful_images):
                doc.add_page_break()

        # Add summary section for failed images
        if failed_images:
            doc.add_page_break()
            doc.add_heading("Processing Errors", level=2)

            error_para = doc.add_paragraph()
            error_para.add_run("The following images encountered processing errors:\n").bold = True
            
            for image_data in failed_images:
                error_para.add_run(f"• Image {image_data['image_number']}: ")
                error_msg = image_data.get('error', 'Unknown error')
                error_para.add_run(f"{error_msg}\n")

        # Add statistics summary
        doc.add_page_break()
        doc.add_heading("Processing Statistics", level=2)
        
        stats_para = doc.add_paragraph()
        stats_para.add_run("Summary Statistics:\n").bold = True
        stats_para.add_run(f"• Total images processed: {results['total_images']}\n")
        stats_para.add_run(f"• Successful images: {len(successful_images)}\n")
        stats_para.add_run(f"• Failed images: {len(failed_images)}\n")
        stats_para.add_run(f"• Success rate: {results.get('success_rate', 0):.1f}%\n")
        
        if images_with_confidence:
            avg_confidence = sum(img['confidence'] for img in images_with_confidence) / len(images_with_confidence)
            stats_para.add_run(f"• Average confidence: {avg_confidence:.1f}%\n")
            
            # Count images by confidence level
            high_conf = len([img for img in images_with_confidence if img['confidence'] >= 85])
            med_conf = len([img for img in images_with_confidence if 70 <= img['confidence'] < 85])
            low_conf = len([img for img in images_with_confidence if img['confidence'] < 70])
            
            stats_para.add_run(f"• High confidence images (≥85%): {high_conf}\n")
            stats_para.add_run(f"• Medium confidence images (70-84%): {med_conf}\n")
            stats_para.add_run(f"• Low confidence images (<70%): {low_conf}\n")

        # Save document
        doc.save(str(doc_path))
        logger.info(f"✓ Tesseract review document created: {doc_path}")

        # Show file info
        file_size_mb = doc_path.stat().st_size / (1024 * 1024)
        logger.info(f"File size: {file_size_mb:.2f} MB")

        return doc_path


def main():
    """Main entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Generate side-by-side review document from Tesseract OCR results'
    )
    parser.add_argument(
        'output_dir',
        help='Path to Tesseract output directory containing results'
    )
    parser.add_argument(
        '-o', '--output',
        default=None,
        help='Custom output path for review document'
    )
    
    args = parser.parse_args()

    try:
        generator = TesseractSideBySideGenerator(args.output_dir)
        doc_path = generator.create_review_document(args.output)

        print("\n🎉 Tesseract review document created successfully!")
        print(f"📄 Open: {doc_path}")
        print("\n💡 Tips:")
        print("- Tesseract works best with printed text")
        print("- Edit the OCR text directly in Word")
        print("- Use Track Changes to see your edits")
        print("- Compare with the original images for accuracy")
        print("- Save frequently as you work")

    except FileNotFoundError as e:
        logger.error(f"File not found: {e}")
        print("\n❓ Make sure you've run the Tesseract processing first:")
        print("1. Install Tesseract OCR")
        print("2. Run the Tesseract image processor")
        print("3. Check that the output directory contains the summary JSON file")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
